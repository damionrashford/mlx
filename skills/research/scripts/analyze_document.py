#!/usr/bin/env python3
"""
Analyze documents from URLs or local paths.
Supports PDF, Word (.docx), text, markdown, and images (with OCR).
Extracted from RivalSearchMCP document_analysis tool.

Usage:
    python3 analyze_document.py <url_or_path> [--max-pages 10] [--no-metadata]
    python3 analyze_document.py https://arxiv.org/pdf/2401.12345.pdf
    python3 analyze_document.py ./papers/paper.pdf --max-pages 20
"""

import argparse
import io
import json
import sys
from datetime import datetime
from pathlib import Path


def check_ocr():
    """Check if OCR dependencies are available."""
    try:
        import importlib.util
        return (
            importlib.util.find_spec("easyocr") is not None
            and importlib.util.find_spec("PIL") is not None
        )
    except Exception:
        return False


def detect_type(path_or_url: str, content_type: str = "") -> str:
    """Detect document type from path/URL and content-type header."""
    s = path_or_url.lower()
    if ".pdf" in s or "application/pdf" in content_type:
        return "pdf"
    elif ".docx" in s or "application/vnd.openxmlformats" in content_type:
        return "docx"
    elif ".txt" in s or "text/plain" in content_type:
        return "txt"
    elif ".md" in s or "text/markdown" in content_type:
        return "markdown"
    elif any(ext in s for ext in [".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"]):
        return "image"
    elif "image/" in content_type:
        return "image"
    return "unknown"


def analyze_pdf(data: bytes, source: str, max_pages: int, extract_metadata: bool) -> dict:
    """Analyze a PDF document."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"status": "error", "error": "pypdf not installed. Run: pip install pypdf"}

    reader = PdfReader(io.BytesIO(data))

    metadata = {}
    if extract_metadata and reader.metadata:
        metadata = {
            "title": reader.metadata.get("/Title", ""),
            "author": reader.metadata.get("/Author", ""),
            "subject": reader.metadata.get("/Subject", ""),
        }

    text_content = []
    pages_to_extract = min(len(reader.pages), max_pages)
    for i in range(pages_to_extract):
        text = reader.pages[i].extract_text()
        if text:
            text_content.append(text)

    full_text = "\n\n".join(text_content)

    return {
        "status": "success",
        "source": source,
        "document_type": "pdf",
        "total_pages": len(reader.pages),
        "pages_extracted": pages_to_extract,
        "metadata": metadata,
        "text": full_text,
        "text_length": len(full_text),
        "timestamp": datetime.now().isoformat(),
    }


def analyze_docx(data: bytes, source: str, extract_metadata: bool) -> dict:
    """Analyze a Word document."""
    try:
        from docx import Document
    except ImportError:
        return {"status": "error", "error": "python-docx not installed. Run: pip install python-docx"}

    document = Document(io.BytesIO(data))
    text_content = [p.text for p in document.paragraphs if p.text.strip()]
    full_text = "\n\n".join(text_content)

    metadata = {}
    if extract_metadata:
        props = document.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
        }

    return {
        "status": "success",
        "source": source,
        "document_type": "docx",
        "paragraphs": len(document.paragraphs),
        "metadata": metadata,
        "text": full_text,
        "text_length": len(full_text),
        "timestamp": datetime.now().isoformat(),
    }


def analyze_text(data: bytes, source: str, doc_type: str) -> dict:
    """Analyze a text-based document."""
    text = None
    for encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        return {"status": "error", "error": "Unable to decode text document"}

    return {
        "status": "success",
        "source": source,
        "document_type": doc_type,
        "text": text,
        "text_length": len(text),
        "lines": len(text.split("\n")),
        "timestamp": datetime.now().isoformat(),
    }


def load_document(source: str) -> tuple:
    """Load document from URL or local path. Returns (bytes, content_type)."""
    if source.startswith("http://") or source.startswith("https://"):
        import requests
        headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"}
        resp = requests.get(source, headers=headers, timeout=60, allow_redirects=True)
        resp.raise_for_status()
        return resp.content, resp.headers.get("content-type", "")
    else:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")
        return path.read_bytes(), ""


def main():
    parser = argparse.ArgumentParser(description="Analyze documents (PDF, Word, text, images)")
    parser.add_argument("source", help="URL or local file path")
    parser.add_argument("--max-pages", type=int, default=10, help="Max pages to extract from PDFs")
    parser.add_argument("--no-metadata", action="store_true", help="Skip metadata extraction")
    parser.add_argument("--json", action="store_true", help="Output raw JSON instead of formatted text")
    args = parser.parse_args()

    try:
        data, content_type = load_document(args.source)
    except Exception as e:
        print(f"Error loading document: {e}", file=sys.stderr)
        sys.exit(1)

    doc_type = detect_type(args.source, content_type)
    extract_metadata = not args.no_metadata

    if doc_type == "pdf":
        result = analyze_pdf(data, args.source, args.max_pages, extract_metadata)
    elif doc_type == "docx":
        result = analyze_docx(data, args.source, extract_metadata)
    elif doc_type in ["txt", "markdown", "unknown"]:
        result = analyze_text(data, args.source, doc_type)
    elif doc_type == "image":
        result = {"status": "error", "error": "Image OCR requires easyocr + pillow. Use pdftotext for PDFs."}
    else:
        result = analyze_text(data, args.source, "unknown")

    if args.json:
        print(json.dumps(result, indent=2))
    else:
        if result["status"] == "error":
            print(f"Error: {result['error']}", file=sys.stderr)
            sys.exit(1)

        if result.get("metadata"):
            meta = result["metadata"]
            if meta.get("title"):
                print(f"Title:  {meta['title']}")
            if meta.get("author"):
                print(f"Author: {meta['author']}")
            print()

        print(f"Type: {result.get('document_type', 'unknown')}")
        if result.get("total_pages"):
            print(f"Pages: {result['pages_extracted']}/{result['total_pages']}")
        print(f"Text length: {result.get('text_length', 0):,} chars")
        print(f"---")
        print(result.get("text", ""))


if __name__ == "__main__":
    main()
